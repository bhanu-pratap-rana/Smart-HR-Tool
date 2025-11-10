"""Document rendering service for generating DOCX and PDF files with branding."""

import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from jinja2 import Template

from backend.app.models.database import CompanyProfile

# WeasyPrint is optional - will use xhtml2pdf as fallback
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as e:
    WEASYPRINT_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning(f"WeasyPrint not available: {e}. Will use xhtml2pdf for PDF generation.")

# xhtml2pdf as fallback for Windows (no system libraries required)
try:
    from xhtml2pdf import pisa
    XHTML2PDF_AVAILABLE = True
except ImportError as e:
    XHTML2PDF_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning(f"xhtml2pdf not available: {e}. PDF generation will be disabled.")

logger = logging.getLogger(__name__)


class DocumentRenderer:
    """Service for rendering professional documents with company branding."""

    def __init__(self, company_profile: Optional[CompanyProfile] = None):
        """
        Initialize document renderer.

        Args:
            company_profile: Company profile for branding (optional)
        """
        self.company = company_profile
        self.templates_dir = Path("backend/app/templates")
        self.templates_dir.mkdir(parents=True, exist_ok=True)

    def render_docx(
        self,
        content: str,
        doc_type: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Render DOCX document with company branding.

        Args:
            content: Markdown content to render
            doc_type: Type of document (e.g., 'job_description', 'offer_letter')
            metadata: Additional metadata (title, date, reference, etc.)

        Returns:
            bytes: DOCX file content
        """
        if metadata is None:
            metadata = {}

        doc = Document()

        # Set document margins (1 inch all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add header with company branding
        self._add_docx_header(doc)

        # Add title
        title_text = metadata.get('title', doc_type.replace('_', ' ').title())
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata (date, reference)
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(
            f"Generated: {metadata.get('date', datetime.utcnow().strftime('%Y-%m-%d'))}\n"
        )
        if metadata.get('reference'):
            metadata_para.add_run(f"Reference: {metadata['reference']}")
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacer

        # Add main content (convert from markdown)
        self._add_markdown_content(doc, content)

        # Add footer with company info
        self._add_docx_footer(doc)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_docx_header(self, doc: Document):
        """Add header with company branding to DOCX."""
        if not self.company:
            return

        header = doc.sections[0].header
        header_para = header.paragraphs[0]

        # Company name
        run = header_para.add_run(f"{self.company.name}\n")
        run.font.size = Pt(14)
        run.font.bold = True

        # Company info
        if self.company.location:
            info_run = header_para.add_run(f"{self.company.location}")
            info_run.font.size = Pt(9)

        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_docx_footer(self, doc: Document):
        """Add footer with company info to DOCX."""
        if not self.company:
            return

        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]

        footer_text = f"© {datetime.utcnow().year} {self.company.name}. All rights reserved."
        if self.company.website:
            footer_text += f"\nWebsite: {self.company.website}"

        footer_para.text = footer_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

    def _add_markdown_content(self, doc: Document, content: str):
        """
        Parse markdown and add to document with proper formatting.

        Supports:
        - Headings (# ## ###)
        - Lists (- * 1.)
        - Bold text (**)
        - Regular paragraphs
        """
        lines = content.split('\n')

        for line in lines:
            line = line.strip()

            if not line:
                continue

            # Headings
            if line.startswith('###'):
                doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)

            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or (line[0].isdigit() and '. ' in line):
                text = line.split('. ', 1)[1] if '. ' in line else line
                doc.add_paragraph(text, style='List Number')

            # Bold text
            elif '**' in line:
                para = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = para.add_run(part)
                    if i % 2 == 1:  # Odd indices are bold
                        run.bold = True

            # Normal paragraph
            else:
                doc.add_paragraph(line)

    def render_pdf(
        self,
        content: str,
        doc_type: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Render PDF document with company branding.
        Uses WeasyPrint if available (better quality), falls back to xhtml2pdf (Windows-compatible).

        Args:
            content: Markdown content to render
            doc_type: Type of document (e.g., 'job_description', 'offer_letter')
            metadata: Additional metadata (title, date, reference, etc.)

        Returns:
            bytes: PDF file content

        Raises:
            RuntimeError: If neither PDF library is available
        """
        if not WEASYPRINT_AVAILABLE and not XHTML2PDF_AVAILABLE:
            raise RuntimeError(
                "No PDF library available. Install either:\n"
                "- WeasyPrint (high quality, requires system libraries): "
                "https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation\n"
                "- xhtml2pdf (Windows-friendly, pure Python): pip install xhtml2pdf"
            )

        if metadata is None:
            metadata = {}

        # Load HTML template
        template_path = self.templates_dir / f"{doc_type}_template.html"

        if template_path.exists():
            with open(template_path) as f:
                template = Template(f.read())
        else:
            # Use default template
            template = Template(self._get_default_html_template())

        # Convert markdown to HTML
        content_html = markdown.markdown(
            content,
            extensions=['extra', 'codehilite', 'tables']
        )

        # Render template
        html_content = template.render(
            company=self.company,
            content=content_html,
            metadata=metadata,
            title=metadata.get('title', doc_type.replace('_', ' ').title()),
            date=metadata.get('date', datetime.utcnow().strftime('%Y-%m-%d')),
            reference=metadata.get('reference', ''),
            year=datetime.utcnow().year
        )

        # Add inline CSS for xhtml2pdf compatibility
        html_with_styles = self._add_inline_styles(html_content)

        # Try WeasyPrint first (better quality)
        if WEASYPRINT_AVAILABLE:
            try:
                pdf_bytes = HTML(string=html_with_styles).write_pdf(
                    stylesheets=[CSS(string=self._get_pdf_styles())]
                )
                logger.info("PDF generated using WeasyPrint")
                return pdf_bytes
            except Exception as e:
                logger.warning(f"WeasyPrint failed: {e}. Trying xhtml2pdf...")

        # Fallback to xhtml2pdf (Windows-compatible)
        if XHTML2PDF_AVAILABLE:
            buffer = BytesIO()
            pisa_status = pisa.CreatePDF(
                html_with_styles,
                dest=buffer
            )

            if pisa_status.err:
                raise RuntimeError(f"xhtml2pdf failed to generate PDF: {pisa_status.err}")

            buffer.seek(0)
            logger.info("PDF generated using xhtml2pdf")
            return buffer.getvalue()

        raise RuntimeError("PDF generation failed with all available libraries")

    def _add_inline_styles(self, html_content: str) -> str:
        """
        Add inline CSS styles to HTML for xhtml2pdf compatibility.
        xhtml2pdf has limited CSS support, so we embed styles directly.
        """
        # Check if HTML already has <style> tag
        if '<style>' in html_content:
            # Insert additional styles after opening <style> tag
            style_css = self._get_xhtml2pdf_styles()
            html_content = html_content.replace('<style>', f'<style>\n{style_css}\n', 1)
        else:
            # Insert <style> tag in <head> if it exists, otherwise before </html>
            style_css = self._get_xhtml2pdf_styles()
            style_tag = f'<style>\n{style_css}\n</style>'

            if '<head>' in html_content:
                html_content = html_content.replace('</head>', f'{style_tag}\n</head>', 1)
            elif '</html>' in html_content:
                html_content = html_content.replace('</html>', f'{style_tag}\n</html>', 1)
            else:
                # Wrap entire content
                html_content = f'<html><head>{style_tag}</head><body>{html_content}</body></html>'

        return html_content

    def _get_xhtml2pdf_styles(self) -> str:
        """
        CSS styles compatible with xhtml2pdf.
        xhtml2pdf supports a subset of CSS 2.1 properties.
        """
        company_name = self.company.name if self.company else "Document"

        return f"""
        body {{
            font-family: Arial, sans-serif;
            font-size: 12pt;
            line-height: 1.6;
            color: #333333;
            margin: 25mm;
        }}

        header {{
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 2px solid #0066cc;
        }}

        header h1 {{
            color: #0066cc;
            font-size: 24pt;
            margin: 0;
        }}

        header .location {{
            color: #666666;
            font-size: 11pt;
            margin-top: 5px;
        }}

        .document-header {{
            text-align: center;
            margin-bottom: 30px;
        }}

        .document-header h2 {{
            color: #333333;
            font-size: 20pt;
            margin: 0;
        }}

        .metadata {{
            color: #666666;
            font-size: 10pt;
            margin-top: 10px;
        }}

        main {{
            margin: 20px 0;
        }}

        h1 {{
            color: #0066cc;
            font-size: 18pt;
            margin-top: 20px;
            margin-bottom: 10px;
        }}

        h2 {{
            color: #0066cc;
            font-size: 16pt;
            margin-top: 15px;
            margin-bottom: 8px;
        }}

        h3 {{
            color: #333333;
            font-size: 14pt;
            margin-top: 12px;
            margin-bottom: 6px;
        }}

        p {{
            margin-bottom: 10px;
            text-align: justify;
        }}

        ul, ol {{
            margin-bottom: 15px;
        }}

        li {{
            margin-bottom: 5px;
        }}

        strong {{
            color: #000000;
            font-weight: bold;
        }}

        footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #dddddd;
            font-size: 9pt;
            color: #666666;
        }}

        table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 15px;
        }}

        th, td {{
            border: 1px solid #dddddd;
            padding: 8px;
            text-align: left;
        }}

        th {{
            background-color: #f2f2f2;
            font-weight: bold;
        }}
        """

    def _get_default_html_template(self) -> str:
        """Default HTML template for documents."""
        return """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ title }}</title>
</head>
<body>
    <header>
        {% if company %}
        <div class="company-info">
            <h1>{{ company.name }}</h1>
            {% if company.location %}
            <p class="location">{{ company.location }}</p>
            {% endif %}
        </div>
        {% endif %}
    </header>

    <div class="document-header">
        <h2>{{ title }}</h2>
        <p class="metadata">
            Generated: {{ date }}
            {% if reference %}
            <br>Reference: {{ reference }}
            {% endif %}
        </p>
    </div>

    <main>
        {{ content|safe }}
    </main>

    <footer>
        {% if company %}
        <p>
            © {{ year }} {{ company.name }}. All rights reserved.
            {% if company.website %}
            <br>Website: {{ company.website }}
            {% endif %}
        </p>
        {% endif %}
    </footer>
</body>
</html>
        """

    def _get_pdf_styles(self) -> str:
        """CSS styles for PDF rendering."""
        company_name = self.company.name if self.company else "Document"

        return f"""
        @page {{
            size: A4;
            margin: 25mm;

            @top-center {{
                content: "{company_name}";
                font-size: 10pt;
                color: #666;
            }}

            @bottom-center {{
                content: "Page " counter(page) " of " counter(pages);
                font-size: 9pt;
                color: #666;
            }}
        }}

        body {{
            font-family: 'Arial', sans-serif;
            font-size: 12pt;
            line-height: 1.6;
            color: #333;
        }}

        header {{
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 2px solid #0066cc;
        }}

        header h1 {{
            color: #0066cc;
            font-size: 24pt;
            margin: 0;
        }}

        header .location {{
            color: #666;
            font-size: 11pt;
            margin-top: 5px;
        }}

        .document-header {{
            text-align: center;
            margin-bottom: 30px;
        }}

        .document-header h2 {{
            color: #333;
            font-size: 20pt;
            margin: 0;
        }}

        .metadata {{
            color: #666;
            font-size: 10pt;
            margin-top: 10px;
        }}

        main {{
            margin: 20px 0;
        }}

        h1 {{
            color: #0066cc;
            font-size: 18pt;
            margin-top: 20px;
            margin-bottom: 10px;
        }}

        h2 {{
            color: #0066cc;
            font-size: 16pt;
            margin-top: 15px;
            margin-bottom: 8px;
        }}

        h3 {{
            color: #333;
            font-size: 14pt;
            margin-top: 12px;
            margin-bottom: 6px;
        }}

        p {{
            margin-bottom: 10px;
            text-align: justify;
        }}

        ul, ol {{
            margin-bottom: 15px;
        }}

        li {{
            margin-bottom: 5px;
        }}

        strong {{
            color: #000;
        }}

        footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            font-size: 9pt;
            color: #666;
        }}
        """


__all__ = ["DocumentRenderer"]
